import os
import sys
from docx import Document
from docx.shared import RGBColor
import re

def extract_text_with_style(paragraph):
    """提取段落文本,保留基本格式"""
    text = ""
    for run in paragraph.runs:
        # 获取文本
        content = run.text
        
        # 处理加粗
        if run.bold:
            content = f"**{content}**"
            
        # 处理斜体
        if run.italic:
            content = f"*{content}*"
            
        # 处理代码样式 (假设使用灰色字体表示代码)
        try:
            if run.font.color.rgb and run.font.color.rgb.is_equivalent_to(RGBColor(128, 128, 128)):
                content = f"`{content}`"
        except:
            pass
            
        text += content
    return text

def convert_docx_to_md(docx_path, md_path):
    """将Word文档转换为Markdown格式"""
    doc = Document(docx_path)
    md_content = []
    
    # 用于跟踪列表状态
    in_list = False
    list_type = None
    list_level = 0
    
    for paragraph in doc.paragraphs:
        text = extract_text_with_style(paragraph)
        
        # 跳过空段落
        if not text.strip():
            if in_list:
                in_list = False
                md_content.append("")
            continue
        
        style_name = paragraph.style.name.lower()
        
        # 处理标题
        if "heading" in style_name:
            level = int(re.search(r"heading (\d+)", style_name).group(1))
            md_content.append(f"\n{'#' * level} {text}\n")
            
        # 处理列表
        elif paragraph.style.name.startswith('List'):
            if not in_list:
                in_list = True
                md_content.append("")
            
            # 检测列表层级
            level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
            indent = "  " * (level - 1)
            
            # 检查是否为有序列表
            if any(text.startswith(f"{i}.") for i in range(1, 10)):
                md_content.append(f"{indent}1. {text.split('.', 1)[1].strip()}")
            else:
                md_content.append(f"{indent}- {text}")
                
        # 处理代码块
        elif paragraph.style.name == 'Code' or "代码" in paragraph.style.name:
            if not any(p.startswith("```") for p in md_content[-3:]):
                md_content.append("\n```c")
            md_content.append(text)
            if len(doc.paragraphs) > doc.paragraphs.index(paragraph) + 1:
                next_style = doc.paragraphs[doc.paragraphs.index(paragraph) + 1].style.name
                if next_style != 'Code' and "代码" not in next_style:
                    md_content.append("```\n")
                    
        # 普通段落
        else:
            if in_list:
                in_list = False
                md_content.append("")
            md_content.append(text + "\n")
    
    # 写入Markdown文件
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))

def process_course_homeworks(course_folder):
    """处理课程文件夹中的所有作业文件"""
    homework_dir = os.path.join(course_folder, 'homeworks')
    if not os.path.exists(homework_dir):
        print(f"作业文件夹不存在: {homework_dir}")
        return
        
    for filename in os.listdir(homework_dir):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            docx_path = os.path.join(homework_dir, filename)
            md_path = os.path.join(homework_dir, filename[:-5] + '.md')
            
            try:
                print(f"正在转换: {filename}")
                convert_docx_to_md(docx_path, md_path)
                print(f"转换完成: {md_path}")
            except Exception as e:
                print(f"转换失败 {filename}: {str(e)}")

def main():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    courses_dir = os.path.join(base_dir, 'courses')
    
    # 遍历所有课程文件夹
    for course in os.listdir(courses_dir):
        course_path = os.path.join(courses_dir, course)
        if os.path.isdir(course_path):
            print(f"\n处理课程: {course}")
            process_course_homeworks(course_path)

if __name__ == "__main__":
    main()